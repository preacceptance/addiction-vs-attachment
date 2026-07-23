#!/usr/bin/env python3
"""3x3 grid on the held-out gold 50: manuals x few-shot sets.

Manuals (docx read verbatim at runtime; embedded tables rendered as text):
  v4          — Coding_Instructions_v4.docx
  v5          — Coding_Instructions_v5.docx        (Bowlby prose + 3 guardrail rules + tables)
  v5simple    — Coding_Instructions_v5_simplified.docx  (same body as v5, NO tables)
Few-shot sets:
  fs30        — legal_fewshot_v3.xlsx
  fs50        — legal_fewshot_gold50.xlsx
  fs80        — legal_fewshot_v3_plus_gold50.xlsx

Output: output/grid_manuals_fewshots_heldout50.xlsx + printed kappa grid.
"""

from __future__ import annotations

import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import pandas as pd
from docx import Document
from openai import OpenAI
from sklearn.metrics import cohen_kappa_score
from tqdm import tqdm

ROOT = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(ROOT / "code"))

import llm_pass_v2 as L  # type: ignore

L.load_dotenv()

MANUALS = {
    "v4":       ROOT / "Coding_Instructions_v4.docx",
    "v5":       ROOT / "Coding_Instructions_v5.docx",
    "v5simple": ROOT / "Coding_Instructions_v5_simplified.docx",
}
FEWSHOTS = {
    "fs30": ROOT / "modified_data" / "legal_fewshot_v3.xlsx",
    "fs50": ROOT / "modified_data" / "legal_fewshot_gold50.xlsx",
    "fs80": ROOT / "modified_data" / "legal_fewshot_v3_plus_gold50.xlsx",
}
OUT = ROOT / "output" / "grid_manuals_fewshots_heldout50.xlsx"

OUTPUT_BLOCK = '''

OUTPUT
  deeper_meaning: one of "addiction", "attachment", "both", "neither"
  reasoning: one sentence justifying the code
  reasoning: one sentence justifying the code'''


def render_table(t) -> str:
    return "\n".join(" | ".join(c.text.strip() for c in row.cells) for row in t.rows)


def build_prompt(docx_path: Path) -> str:
    """Docx body + tables, in document order (tables follow the paragraph stream)."""
    doc = Document(docx_path)
    body = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    for t in doc.tables:
        body += "\n\n" + render_table(t)
    return body + OUTPUT_BLOCK


def main() -> None:
    man = pd.read_csv(ROOT / "code" / "dev" / "_irr_gold50_manifest_legal.csv")
    man["final_code"] = man["final_code"].astype(str).str.strip().str.lower()
    client = OpenAI()

    for m_name, m_path in MANUALS.items():
        L.LEGAL_CODING_SYSTEM = build_prompt(m_path)
        for f_name, f_path in FEWSHOTS.items():
            col = f"llm_{m_name}_{f_name}"
            L.FEW_SHOT_FILE = f_path
            fewshot = L.build_fewshot_messages()

            def task(i):
                comp = L.code_paragraph(client, str(man.at[i, "para_text"]), fewshot)
                return i, comp.deeper_meaning if comp else None

            with ThreadPoolExecutor(max_workers=L.CONCURRENT_WORKERS) as ex:
                futs = [ex.submit(task, i) for i in man.index]
                for fut in tqdm(as_completed(futs), total=len(futs), desc=col):
                    i, code = fut.result()
                    man.at[i, col] = code

    man.to_excel(OUT, index=False)
    print(f"\nSaved {OUT}\n")
    print("=== gold-vs-LLM kappa, held-out 50 ===")
    print(f"{'':10}" + "".join(f"{f:>10}" for f in FEWSHOTS))
    for m_name in MANUALS:
        row = f"{m_name:10}"
        for f_name in FEWSHOTS:
            col = f"llm_{m_name}_{f_name}"
            ok = man[man[col].notna()]
            row += f"{cohen_kappa_score(ok['final_code'], ok[col]):>10.3f}"
        print(row)
    print("\n=== raw agreement ===")
    print(f"{'':10}" + "".join(f"{f:>10}" for f in FEWSHOTS))
    for m_name in MANUALS:
        row = f"{m_name:10}"
        for f_name in FEWSHOTS:
            col = f"llm_{m_name}_{f_name}"
            ok = man[man[col].notna()]
            row += f"{(ok['final_code'] == ok[col]).mean():>10.2f}"
        print(row)


if __name__ == "__main__":
    main()
